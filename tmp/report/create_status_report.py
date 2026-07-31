from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/farhad/Desktop/منظومة القياس ")
OUT = ROOT / "tmp/report/تقرير_حالة_تنفيذ_منظومة_المقياس_2026-07-24.docx"
LOGO = ROOT / "al-miqyas/public/brand/al-amad-logo.png"

NAVY = "0B1026"
PURPLE = "6D35E8"
PURPLE_SOFT = "F2ECFF"
GOLD = "C9A24B"
GREEN = "1D8E5A"
GREEN_SOFT = "EAF7F0"
AMBER = "A96B00"
AMBER_SOFT = "FFF5DF"
RED = "B42318"
RED_SOFT = "FFF0EE"
BLUE = "147EA1"
BLUE_SOFT = "EAF7FB"
GRAY_900 = "1D2433"
GRAY_700 = "495368"
GRAY_500 = "747F94"
GRAY_300 = "D8DDE8"
GRAY_200 = "E8EBF2"
GRAY_100 = "F5F6F9"
WHITE = "FFFFFF"


def rt(text: str) -> str:
    return f"\u202b{text}\u202c"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in edges:
            continue
        tag = f"w:{edge}"
        edge_node = tc_borders.find(qn(tag))
        if edge_node is None:
            edge_node = OxmlElement(tag)
            tc_borders.append(edge_node)
        for key, value in edges[edge].items():
            edge_node.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, color=GRAY_300, size=5):
    spec = {"val": "single", "sz": str(size), "color": color}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=spec,
                bottom=spec,
                start=spec,
                end=spec,
                insideH=spec,
                insideV=spec,
            )


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    keep.set(qn("w:val"), "1")
    p_pr.append(keep)


def set_repeat_header(row):
    set_repeat_table_header(row)


def set_run_font(run, family="Noto Sans Arabic", size=10, bold=False, color=GRAY_900):
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    run._element.rPr.rFonts.set(qn("w:cs"), family)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text, family="Noto Sans Arabic", size=10, bold=False, color=GRAY_900):
    run = paragraph.add_run(text)
    set_run_font(run, family, size, bold, color)
    return run


def add_paragraph(
    doc,
    text="",
    *,
    size=10,
    color=GRAY_900,
    bold=False,
    family="Noto Sans Arabic",
    before=0,
    after=4,
    line=1.15,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    keep=False,
):
    p = doc.add_paragraph()
    set_rtl(p, align)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if keep:
        keep_with_next(p)
    add_run(p, rt(text), family, size, bold, color)
    return p


def add_bullet(doc, text, *, tone="default", size=9.6, level=0):
    colors = {
        "done": GREEN,
        "partial": AMBER,
        "todo": RED,
        "note": BLUE,
        "default": PURPLE,
    }
    marks = {
        "done": "●",
        "partial": "◐",
        "todo": "○",
        "note": "◆",
        "default": "•",
    }
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.right_indent = Cm(0.25 + level * 0.35)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.08
    add_run(p, marks[tone] + " ", "Noto Sans Symbols 2", size, True, colors[tone])
    add_run(p, rt(text), "Noto Sans Arabic", size, False, GRAY_900)
    return p


def add_section_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, rt(text), "Noto Kufi Arabic", 15, True, NAVY)
    if subtitle:
        s = add_paragraph(doc, subtitle, size=9.3, color=GRAY_500, after=7, keep=True)
        s.paragraph_format.keep_with_next = True
    return p


def add_subhead(doc, text, tone=PURPLE):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    add_run(p, rt(text), "Noto Kufi Arabic", 10.5, True, tone)
    return p


def set_cell_text(cell, text, *, size=8.6, bold=False, color=GRAY_900, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p, align)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    add_run(p, rt(text), "Noto Sans Arabic", size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_callout(doc, title, body, *, fill=PURPLE_SOFT, accent=PURPLE, icon=""):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.0)
    table.columns[1].width = Cm(0.55)
    main = table.cell(0, 0)
    rail = table.cell(0, 1)
    set_cell_shading(main, fill)
    set_cell_shading(rail, accent)
    set_cell_border(main, top={"val": "single", "sz": "4", "color": accent},
                    bottom={"val": "single", "sz": "4", "color": accent},
                    start={"val": "single", "sz": "4", "color": accent},
                    end={"val": "nil"})
    set_cell_border(rail, top={"val": "single", "sz": "4", "color": accent},
                    bottom={"val": "single", "sz": "4", "color": accent},
                    start={"val": "nil"}, end={"val": "single", "sz": "4", "color": accent})
    set_cell_margins(main, 130, 160, 130, 160)
    set_cell_margins(rail, 0, 0, 0, 0)
    p = main.paragraphs[0]
    set_rtl(p)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, rt(f"{icon} {title}".strip()), "Noto Kufi Arabic", 9.8, True, accent)
    p2 = main.add_paragraph()
    set_rtl(p2)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    add_run(p2, rt(body), "Noto Sans Arabic", 9.1, False, GRAY_900)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, "Noto Sans Arabic", 8, False, GRAY_500)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans Arabic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Arabic")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Arabic")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(GRAY_900)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(18.3))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(14.8)
    table.columns[1].width = Cm(3.5)
    set_cell_text(table.cell(0, 0), "منظومة المقياس · تقرير حالة التنفيذ المحدث", size=8, bold=True, color=NAVY)
    set_cell_text(table.cell(0, 1), "24 يوليو 2026", size=8, color=GRAY_500, align=WD_ALIGN_PARAGRAPH.LEFT)
    for cell in table.row_cells(0):
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": PURPLE})

    footer = section.footer
    p = footer.paragraphs[0]
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(2)
    add_run(p, rt("شركة الأمد التقنية | منظومة المقياس | "), "Noto Sans Arabic", 8, False, GRAY_500)
    add_page_number(p)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(4.6))

    p = doc.add_paragraph()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, rt("منظومة المقياس"), "Noto Kufi Arabic", 24, True, NAVY)

    p = doc.add_paragraph()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, rt("تقرير حالة التنفيذ المحدث"), "Noto Kufi Arabic", 20, True, PURPLE)

    p = doc.add_paragraph()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(18)
    add_run(p, rt("مراجعة فعلية لما أُنجز وما تبقى، مبنية على الملفات والكود والمهاجرات والفحص التنفيذي"), "Noto Sans Arabic", 11, False, GRAY_700)

    add_callout(
        doc,
        "الحكم التنفيذي",
        "المشروع ليس جاهزًا للإطلاق الإنتاجي. توجد قاعدة بيانات قوية وتدفقات حقيقية للبرامج والدفعات والمتدربين، لكن لا توجد حتى الآن رحلة كاملة من القبلي إلى xAPI ثم البعدي والتقرير والشهادة. أجزاء رئيسية من الواجهة ما زالت بيانات ثابتة ومحاكاة.",
        fill=RED_SOFT,
        accent=RED,
        icon="",
    )

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(11.8)
    table.columns[1].width = Cm(5.4)
    rows = [
        ("تاريخ لقطة المراجعة", "24 يوليو 2026"),
        ("النطاق", "الملفات الدلالية للمشروع مع استبعاد الاعتماديات والمخرجات المولدة"),
        ("مرجع القياس", "قائمة مهام التنفيذ الشاملة، الإصدار 1.0"),
        ("الإصدار", "تقرير تدقيق الحالة 2.0"),
    ]
    for idx, (k, v) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), v, size=9.2)
        set_cell_text(table.cell(idx, 1), k, size=9.2, bold=True, color=PURPLE)
        set_cell_shading(table.cell(idx, 1), PURPLE_SOFT)
    set_table_borders(table)

    add_paragraph(doc, "شركة الأمد التقنية · مكة المكرمة", size=10, color=GRAY_500, before=34, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


TASKS = [
    {
        "n": 0,
        "title": "حسم النطاق والقرارات قبل البرمجة",
        "status": "جزئي موثق",
        "tone": "partial",
        "summary": "تم حسم عدد من القرارات المعمارية في سجل مستقل، لكن الملكيات والمسؤوليات ومؤشرات النجاح وخطة المخاطر والرجوع لم تُغلق.",
        "done": [
            "تعريف المنتج كنظام قياس وإثبات، وليس منصة XR، موثق في ملفات الحزمة.",
            "حسم مسار المتدرب على ‎/t/[traineeCode]‎، ومسار Jotform عبر Next.js Route Handler، ونموذج cohort_reports المنفصل.",
            "اعتماد أن شرط الشهادة هو نتيجة الاختبار البعدي 80% أو أكثر، وأن الثقة والأداء اللحظي أدلة أثر لا تمنع الشهادة.",
            "اعتماد رقم شهادة تسلسلي مع verify_code عشوائي، ومنع الحذف التلقائي الصامت.",
        ],
        "partial": [
            "قرار البيئة الواحدة موثق، لكنه يخالف قائمة المهام التي تطلب Development وStaging وProduction؛ يلزم استثناء معتمد أو تعديل القرار.",
            "قاموس xAPI وشروط البيانات الناقصة لم تُعتمد مع فريق AmadXR.",
        ],
        "todo": [
            "تعيين Product Owner والمالك التقني ومسؤولي قاعدة البيانات والتكامل والاختبار بأسماء فعلية.",
            "اعتماد سياسة PDPL للاحتفاظ والحذف والتصدير.",
            "إنشاء Risk Register وChange Log وخطة Rollback قابلة للتنفيذ.",
            "تثبيت مؤشرات 98% للمطابقة وزمن التقرير أقل من دقيقة ضمن خطة اختبار.",
        ],
        "evidence": "ChatGPT/PROJECT_DECISIONS.md، docs/01-project-brief.md، docs/02-architecture.md",
        "acceptance": "غير مغلقة؛ توجد قرارات حرجة ومسؤوليات تشغيلية بلا اعتماد.",
    },
    {
        "n": 1,
        "title": "الحسابات والبنية التحتية المملوكة للأمد",
        "status": "جزئي وغير متحقق خارجيًا",
        "tone": "partial",
        "summary": "مشروع Supabase فعلي موجود ومربوط محليًا، لكن ملكية GitHub وVercel والدومين و2FA والاستعادة لا يمكن إثباتها من الملفات.",
        "done": [
            "وجود مشروع Supabase مستقل ومفاتيح عامة محلية مسماة دون كشف القيم.",
            "ملف .gitignore يمنع .env، ولم تُطبع أو تُضمّن أسرار في هذا التقرير.",
            "الدالة العامة get_public_trainee_route استجابت من Supabase فعليًا بحالة HTTP 200.",
        ],
        "partial": [
            "ملف .env.example موجود فقط في حزمة التسليم القديمة ويستخدم NEXT_PUBLIC_SUPABASE_ANON_KEY، بينما التطبيق الحالي يستخدم NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY.",
            "لا يوجد NEXT_PUBLIC_APP_URL في البيئة الحالية، ما يمنع اختبار QR من هاتف آخر عبر عنوان loopback.",
        ],
        "todo": [
            "إثبات ملكية GitHub وSupabase وVercel والدومين للأمد، وتوثيق المالكين والاستعادة و2FA.",
            "إنشاء .env.example داخل al-miqyas متوافق مع الأسماء الحالية.",
            "حسم استراتيجية البيئات، وضبط الدومين وSSL واختبار Safari على هاتف فعلي.",
            "إضافة فحص تسرب أسرار وميزانيات وتنبيهات استهلاك.",
        ],
        "evidence": "al-miqyas/.env.local، al-miqyas/.gitignore، al-miqyas-handoff/.env.example، فحص Supabase قراءة فقط",
        "acceptance": "غير مغلقة؛ الملكية والدومين و2FA لم تُثبت.",
    },
    {
        "n": 2,
        "title": "تهيئة المستودع ومعايير التطوير",
        "status": "جزئي حقيقي",
        "tone": "partial",
        "summary": "التطبيق مبني فعليًا باستخدام Next.js وTypeScript وRTL، والبناء ينجح، لكن لا يوجد CI أو lint أو test أو مستودع Git صالح محليًا.",
        "done": [
            "Next.js 16.2.10 مع App Router وReact 19 وTypeScript strict.",
            "Tailwind CSS 4، اتجاه RTL، خصائص CSS منطقية، وخطوط IBM Plex.",
            "أوامر dev وbuild وstart وtypecheck، ونجاح typecheck وbuild بتاريخ المراجعة.",
            "إعداد عميل Supabase للمتصفح وعميل خادم داخل proxy لتحديث الجلسة.",
        ],
        "partial": [
            "الوصول إلى Supabase ومنطق العرض موجودان داخل مكونات React كبيرة؛ الفصل الطبقي محدود.",
            "نسخ قديمة غير مستخدمة ما زالت في المشروع: globals 2.css وicons 2.tsx وtrainee-workspace بنسختين.",
        ],
        "todo": [
            "إضافة ESLint وPrettier وأمر lint وأمر test.",
            "إنشاء GitHub Actions وحماية main وقوالب PR وIssue.",
            "إضافة README فعلي داخل al-miqyas يتيح تشغيل المشروع دون معرفة شفوية.",
            "إضافة أنواع قاعدة بيانات مولدة وطبقة أخطاء واستجابات وسجلات موحدة.",
            "تنظيف الملفات القديمة بعد التحقق من عدم استخدامها.",
        ],
        "evidence": "al-miqyas/package.json، tsconfig.json، next.config.ts، app/layout.tsx، نتيجة build وtypecheck",
        "acceptance": "غير مغلقة؛ CI وlint وtest وREADME وحماية المستودع غائبة.",
    },
    {
        "n": 3,
        "title": "تصميم قاعدة البيانات وتنفيذ المهاجرات",
        "status": "متقدم لكنه غير مغلق",
        "tone": "partial",
        "summary": "هذا أقوى جزء في المشروع: 18 جدولًا و34 سياسة RLS و13 دالة SQL عبر 7 مهاجرات. مع ذلك لا توجد اختبارات آلية أو ERD أو Data Dictionary أو إثبات استعادة.",
        "done": [
            "إضافة organizations وmemberships وprograms وprogram_versions وjotform_forms وcohorts وtrainees وenrollments.",
            "إضافة webhook_ingestions وassessments وorg_api_keys وxapi_statements وimpact_reports وcohort_reports وcertificates وaudit_logs.",
            "قيود مركبة تمنع الربط بين الجهات، وفهارس وUNIQUE لـ submission_id وstatement_id والرموز.",
            "RLS مفعل على الجداول الثمانية عشر، وسياسات viewer للقراءة فقط.",
            "حماية نسخ البرامج المنشورة، وتجهيز snapshots للتقارير والشهادات.",
            "المهاجرات 001 إلى 007 موثقة في التسليم على أنها طُبقت يدويًا بنجاح.",
        ],
        "partial": [
            "لا توجد سياسات DELETE لمعظم بيانات الأعمال؛ المنع الافتراضي آمن، لكن إجراءات الأرشفة أو الحذف الإداري غير مكتملة.",
            "توليد AMD-XXXXX يبحث عن رمز فريد، لكن تصادم سباق متزامن نادر قد يفشل الإدراج دون إعادة محاولة للمعاملة.",
        ],
        "todo": [
            "إصدار ERD نهائي وData Dictionary.",
            "اختبار المهاجرات على قاعدة فارغة وعلى Staging ببيانات فعلية.",
            "إضافة Seed آمن واختبارات القيود والفهارس ومنع التكرار.",
            "توثيق النسخ الاحتياطي والاستعادة وRollback قبل أي مهاجرة إنتاجية جديدة.",
            "عدم تشغيل migration 008 قبل اعتماد نماذج Jotform ومفتاح الإجابة.",
        ],
        "evidence": "al-miqyas/supabase/migrations/202607230001_foundation.sql حتى 202607230007_public_trainee_route.sql",
        "acceptance": "غير مغلقة؛ التصميم قوي، لكن بوابة الوثائق والاختبارات والاستعادة غير موجودة.",
    },
    {
        "n": 4,
        "title": "تسجيل الدخول وتعدد الجهات والصلاحيات",
        "status": "جزئي حقيقي",
        "tone": "partial",
        "summary": "تسجيل الدخول والحماية وRLS حقيقية، لكن التطبيق يستخدم كلمة مرور لا OTP، واختيار الجهة في الغلاف ثابت ولا يتبع عضويات المستخدم.",
        "done": [
            "تسجيل دخول فعلي عبر Supabase signInWithPassword وتسجيل خروج فعلي.",
            "proxy يحمي المسارات الخاصة ويعيد غير المسجل إلى /login.",
            "أدوار owner وtrainer وviewer وplatform_admin ممثلة في قاعدة البيانات.",
            "RLS يمنع viewer من الكتابة في مسارات الأعمال الرئيسية.",
        ],
        "partial": [
            "الشاشات الحقيقية تختار أول عضوية نشطة بدل سياق جهة صريح عند تعدد العضويات.",
            "تجديد الجلسة عبر proxy موجود، لكن تجربة انتهاء الجلسة والتعطيل والإزالة غير مكتملة.",
            "platform_admin موجود بنيويًا، ولا توجد واجهة إدارة شاملة مثبتة.",
        ],
        "todo": [
            "حسم كلمة المرور مقابل OTP وتحديث الوثائق والواجهة وفق قرار واحد.",
            "ربط مبدل الجهة بالعضويات الفعلية وإزالة أسماء المستخدم والجهات الثابتة.",
            "إنشاء إدارة مستخدمين وتعطيل عضوية مع حفظ التاريخ.",
            "كتابة RLS Tests لحسابين في جهتين، واختبار URL وAPI وService Role.",
        ],
        "evidence": "components/auth-pages.tsx، proxy.ts، migrations/202607230002_rls_policies.sql، components/app-shell.tsx",
        "acceptance": "غير مغلقة؛ عزل RLS مصمم لكن غير مختبر آليًا وسياق الجهة في الغلاف تجريبي.",
    },
    {
        "n": 5,
        "title": "إدارة الجهات والهوية البصرية",
        "status": "هيكل قاعدة وواجهة محاكاة",
        "tone": "todo",
        "summary": "الجداول والهوية العامة موجودة، لكن إدارة الجهة والأعضاء والدعوات والشعارات ليست تدفقًا حقيقيًا.",
        "done": [
            "جداول organizations وmemberships وmembership_invitations موجودة مع RLS.",
            "أصول شعار الأمد موجودة ويستخدم التطبيق الشعار الفعلي.",
            "Design Tokens للألوان موجودة في globals.css.",
        ],
        "partial": [
            "app-shell يعرض جهتين ثابتتين ومبدلًا وهميًا وبيانات مشغل ثابتة.",
            "brand-spec.md ما زال يقول إن الشعار الرسمي pending رغم وجود أصول فعلية؛ الوثيقة قديمة.",
        ],
        "todo": [
            "صفحة حقيقية لإنشاء وتعديل الجهة والتحقق من slug واللون والشعار والحالة.",
            "Supabase Storage وسياسات رفع الشعارات.",
            "دعوة أعضاء وتغيير أدوار وإلغاء عضوية ومنع إزالة آخر owner.",
            "ربط الهوية الفعلية بالتقارير والشهادات واختبار جهتين.",
        ],
        "evidence": "components/app-shell.tsx، components/dashboard-pages.tsx، brand-spec.md، public/brand، الجداول والسياسات",
        "acceptance": "غير مغلقة؛ لا يمكن إدارة جهتين وأعضائهما وهويتهما فعليًا من الواجهة.",
    },
    {
        "n": 6,
        "title": "البرامج والدفعات والتسجيلات",
        "status": "جزئي حقيقي",
        "tone": "partial",
        "summary": "إنشاء وتعديل المسودات والدفعات وتسجيل المتدرب تعمل فعليًا، لكن النشر ودورة الحالات والربط بالتكاملات غير مكتملة.",
        "done": [
            "إنشاء برنامج مع نسخة أولى عبر RPC ذري، وتعديل المسودة مع Audit Log.",
            "إنشاء دفعة وتعديلها، ومنع تشغيلها قبل نشر البرنامج وربط Jotform وxAPI.",
            "إنشاء متدرب وتسجيله في دفعة عبر RPC ذري مع منع التكرار.",
            "قيود نسخ البرامج تمنع تعديل النسخة المنشورة بصمت.",
        ],
        "partial": [
            "حالات الدفعة والبرنامج موجودة في القاعدة، لكن واجهة النشر والتفعيل والإغلاق غير مكتملة.",
            "قواعد المعرفة الأساسية محفوظة، لكن answer_key ونماذج Jotform وقاموس xAPI غير موصولة.",
            "تقدم المتدرب مشتق من البيانات المتاحة، لكنه لن يكتمل قبل وصول التقييمات والأحداث.",
        ],
        "todo": [
            "اعتماد وإنشاء نسخ البرامج ونشرها وتفعيلها من الواجهة.",
            "ربط النموذج القبلي والبعدي ومفتاح الإجابة وقاموس xAPI بالإصدار.",
            "إدارة انتقالات enrollment حتى reported وcertified.",
            "واجهة تفاصيل دفعة حقيقية بدل غرفة الدفعة التجريبية.",
        ],
        "evidence": "components/programs-page.tsx، components/trainees-page.tsx، migrations 003 إلى 006",
        "acceptance": "غير مغلقة؛ الإنشاء الحقيقي يعمل، لكن دورة التشغيل من نشر البرنامج إلى اكتمال الرحلة لا تعمل.",
    },
    {
        "n": 7,
        "title": "المتدرب والمعرف الموحد وQR",
        "status": "جزئي حقيقي",
        "tone": "partial",
        "summary": "AMD-XXXXX والتسجيل والتوجيه والـQR حقيقية، لكن البطاقة والطباعة واختبارات الهاتف وسلامة تعدد التسجيلات غير مكتملة.",
        "done": [
            "توليد AMD-XXXXX في Supabase بأبجدية آمنة وقيد UNIQUE.",
            "الرمز غير قابل للتعديل من واجهة المتدرب، مع Audit Log لتعديل الملف الشخصي فقط.",
            "البحث بالرمز والاسم والجوال، وصفحة متدرب حقيقية تعرض الأدلة الموجودة فقط.",
            "QR فعلي يولد في المتصفح، وصفحة /t/[traineeCode] تستدعي RPC عامة دون عرض بيانات شخصية.",
            "التوجيه يدعم فتح نماذج Jotform مع prefill عند ربطها.",
        ],
        "partial": [
            "المطلوب توليد QR من الخادم، بينما التنفيذ الحالي يتم في المتصفح.",
            "الدالة العامة تختار أحدث تسجيل للرمز؛ هذا لا يحقق شرط عدم الاعتماد على الرمز وحده عند تعدد البرامج.",
            "المسار العام يعتمد على رمز قابل للتخمين نسبيًا دون Rate Limiting أو طبقة API وسيطة.",
        ],
        "todo": [
            "إنشاء بطاقة قابلة للطباعة والحفظ بهوية الجهة.",
            "تعريف سياسة منع التكرار والموافقة وفق PDPL واستيراد CSV إن لزم.",
            "ربط التوجيه بتسجيل أو سياق محدد عند تعدد البرامج.",
            "إضافة Rate Limiting واختبارات iPhone وAndroid وإضاءة وطباعة.",
        ],
        "evidence": "generate_trainee_code، create_trainee_with_enrollment، trainee-details-page.tsx، trainee-routing-page.tsx، RPC العامة",
        "acceptance": "غير مغلقة؛ المعرف والتوجيه يعملان، لكن البطاقة والاختبارات وسياق التسجيل المتعدد غير مغلقة.",
    },
    {
        "n": 8,
        "title": "تكامل Jotform",
        "status": "لم يبدأ تشغيليًا",
        "tone": "todo",
        "summary": "قاعدة البيانات ومسار التوجيه جاهزان لاستقبال التكامل، لكن لا توجد نماذج معتمدة أو endpoint أو webhook أو تصحيح أو reconciliation.",
        "done": [
            "جدول jotform_forms وجدول webhook_ingestions وجدول assessments جاهزة بنيويًا.",
            "التوجيه العام يعرف بناء رابط prefill عندما توجد form_id وtrainee_field_name.",
            "submission_id عليه قيد فريد في assessments.",
        ],
        "partial": [
            "القرار المعماري حسم Next.js Route Handler، لكن app/api غير موجود أصلًا.",
            "ملف البيئة القديم يسرد JOTFORM_API_KEY وJOTFORM_WEBHOOK_SECRET، ولا توجد استخدامات برمجية لهما.",
        ],
        "todo": [
            "اعتماد النموذج القبلي والبعدي وحقل traineeId المقفل ومفتاح الإجابة.",
            "تنفيذ endpoint موحد وآمن، والتحقق من المصدر، وحفظ raw payload المحمي.",
            "استخراج submission_id والإجابات والتوقيت والنوع والجهة والتسجيل.",
            "التصحيح بإصدار البرنامج، وحفظ الدرجة والثقة والأصل.",
            "Reconciliation دوري يستخدم نفس مسار المعالجة ويمنع التكرار.",
            "لوحة للفشل وعدم الارتباط واختبارات جوال وفشل Webhook.",
        ],
        "evidence": "لا يوجد app/api؛ الجداول موجودة في foundation.sql؛ PROJECT_DECISIONS.md",
        "acceptance": "غير مغلقة؛ لا يوجد إرسال Jotform واحد يُعالج فعليًا.",
    },
    {
        "n": 9,
        "title": "تكامل xAPI والأداء اللحظي",
        "status": "لم يبدأ تشغيليًا",
        "tone": "todo",
        "summary": "الجداول والقيود جاهزة، لكن لا يوجد POST /api/xapi ولا مفاتيح تشغيل ولا تحقق ولا بث حي حقيقي.",
        "done": [
            "جداول org_api_keys وxapi_statements مع key_hash وstatement_id فريد.",
            "قيود تربط الحدث بالجهة والتسجيل والبرنامج.",
            "واجهة تجريبية توضح الشكل المطلوب للأحداث والمقاييس.",
        ],
        "partial": [
            "شاشة الجلسة الحية الحالية محاكاة ببيانات ثابتة وعدّادات محلية.",
            "قاموس مبدئي للأفعال موجود في النموذج والسيناريو، لكنه غير معتمد أو منفذ.",
        ],
        "todo": [
            "إنشاء POST /api/xapi لمفرد أو مصفوفة مع Bearer key وRate Limiting.",
            "إنشاء وتدوير وإلغاء مفاتيح الجهات مع عرضها مرة واحدة.",
            "التحقق من actor.account.name وregistration وprogramId والجهة والوحدات والتوقيت.",
            "Idempotency ورفض قابل للتتبع وربط الجلسة والتسجيل.",
            "اعتماد قاموس xAPI مع AmadXR واختبار جلسة فعلية وانقطاع الشبكة.",
            "استبدال شاشة المحاكاة ببيانات مباشرة.",
        ],
        "evidence": "foundation.sql، dashboard-pages.tsx، cohort-room.tsx، al-miqyas-v2.jsx، غياب app/api",
        "acceptance": "غير مغلقة؛ لا يمكن استقبال statement واحد في التطبيق الحالي.",
    },
    {
        "n": 10,
        "title": "محرك التصحيح والقياس والأثر",
        "status": "هيكل بيانات فقط",
        "tone": "todo",
        "summary": "لا يوجد scoring.ts أو محرك أثر أو اختبارات حساب. الأرقام الصحيحة تظهر في النموذج التجريبي فقط.",
        "done": [
            "جداول impact_reports وcohort_reports تخزن metrics_snapshot وrule_version.",
            "قرار الشهادة المعرفي 80% موثق.",
            "الأمثلة المرجعية min 5→8 وSD 1.62→0.79 موثقة في الملفات.",
        ],
        "partial": [
            "صفحة تفاصيل المتدرب تعرض assessments وxAPI وimpact_reports إن وجدت، لكنها لا تحسبها.",
        ],
        "todo": [
            "بناء scoring.ts أو طبقة خادم مكافئة مع answer_key بإصدار.",
            "حساب Knowledge Delta وConfidence Delta وملخص الأداء اللحظي.",
            "بناء التقرير الفردي والجماعي مع sample matching وCeiling Effect.",
            "سياسة البيانات الناقصة وإعادة الحساب دون تغيير الشهادات بصمت.",
            "Unit Tests للحالات المرجعية والقسمة على صفر والعينات الصغيرة والقيم المفقودة.",
        ],
        "evidence": "لا يوجد scoring.ts؛ foundation.sql؛ trainee-details-page.tsx؛ al-miqyas-v2.jsx",
        "acceptance": "غير مغلقة؛ لا يمكن إعادة إنتاج قرار أو تقرير من بيانات فعلية.",
    },
    {
        "n": 11,
        "title": "الواجهات ولوحات التحكم",
        "status": "مختلط: حقيقي ومحاكاة",
        "tone": "partial",
        "summary": "نظام التصميم متماسك وRTL متجاوب، لكن غالبية الشاشات التشغيلية ما زالت محاكاة. هذا الخلط يمثل مخاطرة تضليل قبل الإطلاق.",
        "done": [
            "تصميم عربي RTL، Design Tokens، حالات focus، responsive breakpoints، وprefers-reduced-motion.",
            "شاشة دخول حقيقية، برامج حقيقية، متدربون حقيقيون، تفاصيل متدرب حقيقية، ومسار عام حقيقي.",
            "نجاح إنتاج 15 مسارًا في build.",
        ],
        "partial": [
            "Loading وError موجودة في الصفحات الحقيقية، لكن ليست موحدة في كل الشاشات.",
            "الصلاحيات تخفي بعض الأزرار حسب role، لكن الغلاف نفسه يعرض بيانات ثابتة.",
            "register يستقبل طلبًا تجريبيًا ولا ينشئ حسابًا أو دعوة.",
        ],
        "todo": [
            "استبدال dashboard وsessions وreports وcertificates وorganizations وsettings وaccount وcohort room ببيانات حقيقية.",
            "إزالة أو وسم المحاكاة بوضوح حتى لا تختلط ببيانات الإنتاج.",
            "إكمال حالات Loading وEmpty وError وSuccess والتأكيدات لكل فعل حساس.",
            "اختبارات Safari وChrome والجوال والنصوص الطويلة والتباين ولوحة المفاتيح.",
        ],
        "evidence": "globals.css، app routes، programs-page.tsx، trainees-page.tsx، dashboard-pages.tsx، app-shell.tsx",
        "acceptance": "غير مغلقة؛ شرط القراءة من بيانات حقيقية غير متحقق لمعظم الواجهة.",
    },
    {
        "n": 12,
        "title": "التقارير الفردية والجماعية",
        "status": "محاكاة مع جداول جاهزة",
        "tone": "todo",
        "summary": "التقارير تعرض القصة الصحيحة بصريًا، لكنها لا تعتمد على محرك حساب أو بيانات فعلية.",
        "done": [
            "تصميم تقرير فردي وجماعي يبرز رفع الأضعف وضغط التشتت.",
            "جداول impact_reports وcohort_reports تدعم snapshots والإصدارات.",
            "التصميم يذكر اختلاف حجم العينة 12 قبلي و7 بعدي في النموذج التجريبي.",
        ],
        "partial": [
            "زر Export PDF في التقرير التجريبي ينفذ window.print فقط.",
            "صفحة تفاصيل المتدرب تعرض وجود المصادر أو فقدها، لكنها ليست تقرير أثر محسوبًا.",
        ],
        "todo": [
            "توليد تقارير من محرك القياس الفعلي.",
            "مرشحات البرنامج والدفعة والتاريخ وصلاحيات العرض والتنزيل.",
            "رقم إصدار وتاريخ الحساب ومصدر البيانات.",
            "طباعة نظيفة واختبار عدم فتح تقرير جهة أخرى.",
        ],
        "evidence": "dashboard-pages.tsx، trainee-details-page.tsx، impact_reports وcohort_reports",
        "acceptance": "غير مغلقة؛ لا يوجد تقرير فعلي قابل لإعادة الحساب والمراجعة.",
    },
    {
        "n": 13,
        "title": "الشهادات والتحقق",
        "status": "هيكل قوي وواجهة محاكاة",
        "tone": "todo",
        "summary": "جدول الشهادات صُمم جيدًا، لكن الإصدار والإلغاء والقالب وQR والتحقق العام ليست تدفقات حقيقية.",
        "done": [
            "certificate_serial وcertificate_number وverify_code فريد وعشوائي بنيويًا.",
            "حالات valid وrevoked وsuperseded، وسبب وتاريخ الإلغاء، وsnapshot.",
            "قيد يمنع أكثر من شهادة valid للتسجيل نفسه.",
        ],
        "partial": [
            "صفحة /verify/[verifyCode] عامة لكنها بيانات ثابتة وتبديل محلي بين valid وrevoked وnotfound.",
            "صفحة الشهادات وزر الإصدار محاكاة، ولا يوجد RPC أو endpoint للإصدار.",
        ],
        "todo": [
            "محرك إصدار مقيد بالبعدي 80% وباكتمال البيانات الإلزامية.",
            "RPC أو Route Handler للإصدار والإلغاء وإعادة الإصدار مع Audit Log.",
            "قالب شهادة فعلي بهوية الجهة وQR إلى /verify/{code}.",
            "صفحة تحقق تقرأ الحد الأدنى الضروري فقط ولا تعرض الدرجات.",
            "اختبار شهادة مطبوعة على Safari وحالات الإلغاء والرمز غير الموجود.",
        ],
        "evidence": "foundation.sql، public-pages.tsx، dashboard-pages.tsx، PROJECT_DECISIONS.md",
        "acceptance": "غير مغلقة؛ لا يمكن إصدار أو إلغاء أو التحقق من شهادة حقيقية.",
    },
    {
        "n": 14,
        "title": "الأمان والخصوصية والامتثال",
        "status": "أساس جزئي",
        "tone": "partial",
        "summary": "RLS والقيود والأسرار المحلية نقطة جيدة، لكن PDPL والحماية الطرفية والنسخ الاحتياطي والاستجابة للحوادث غير مكتملة.",
        "done": [
            "RLS مفعل على 18 جدولًا، وقيود cross-tenant قوية، وviewer بلا كتابة.",
            ".env.local مستبعد من Git، ولا توجد مفاتيح خادم مستخدمة في الواجهة.",
            "الصفحة العامة للمتدرب لا تعرض اسمًا أو جوالًا أو بريدًا.",
        ],
        "partial": [
            "الدالة العامة تكشف حالة تشغيل محدودة عبر رمز AMD-XXXXX دون Rate Limiting.",
            "التحقق من المدخلات موجود داخل RPCs، لكن لا توجد طبقة endpoint موحدة أو schema validation.",
        ],
        "todo": [
            "جرد البيانات الشخصية وإشعار الخصوصية ومدة الاحتفاظ وآليات الطلب والتصحيح والحذف وفق PDPL.",
            "Rate Limiting وSecurity Headers وCSP وCSRF حسب endpoints.",
            "فحص الاعتماديات والثغرات ومراجعة RLS مستقلة.",
            "نسخ احتياطي واختبار استعادة فعلي وIncident Response Runbook وتدوير مفاتيح.",
        ],
        "evidence": "مهاجرات foundation وRLS، .gitignore، proxy.ts، get_public_trainee_route",
        "acceptance": "غير مغلقة؛ لا توجد مراجعة أمنية مستقلة أو اختبار استعادة أو سياسة PDPL معتمدة.",
    },
    {
        "n": 15,
        "title": "المراقبة والتشغيل والدعم",
        "status": "بداية محدودة",
        "tone": "todo",
        "summary": "Audit Log موجود لبعض العمليات، لكن لا توجد مراقبة أو تنبيهات أو tracing أو إعادة معالجة.",
        "done": [
            "جدول audit_logs وسياسات قراءة owner.",
            "RPCs الحالية تسجل إنشاء وتعديل البرامج والدفعات والمتدربين.",
            "webhook_ingestions يتضمن حالات ومحاولات ورسالة خطأ بنيويًا.",
        ],
        "partial": [
            "لا توجد شاشة حقيقية للسجلات أو الفشل؛ المعروض في الجلسات محاكاة.",
        ],
        "todo": [
            "Correlation ID وتسجيل نجاح وفشل Jotform وxAPI دون بيانات حساسة.",
            "Uptime وError Tracking وتنبيهات توقف reconciliation وأخطاء xAPI.",
            "صفحة للأحداث المرفوضة والإرسالات غير المرتبطة.",
            "إعادة معالجة آمنة وتدوير مفاتيح وتصحيح ارتباط عالي الصلاحية مع Audit Log.",
            "إجراءات تشغيل يومية وأسبوعية.",
        ],
        "evidence": "audit_logs وwebhook_ingestions، RPCs 003 إلى 006، sessions التجريبية",
        "acceptance": "غير مغلقة؛ لا يمكن تتبع إرسال واحد أو تنبيه فريق التشغيل عند فشله.",
    },
    {
        "n": 16,
        "title": "الاختبارات وضمان الجودة",
        "status": "غير منفذ كمنظومة QA",
        "tone": "todo",
        "summary": "نجح build وtypecheck فقط. لا يوجد أي ملف اختبار أو إطار اختبار أو CI.",
        "done": [
            "نجاح tsc --noEmit بتاريخ 24 يوليو 2026.",
            "نجاح next build --webpack وإنشاء 15 مسارًا.",
            "نجاح فحص محلي: /login يعيد 200، و/dashboard يعيد 307 إلى login، والمسارات العامة تعيد 200.",
        ],
        "partial": [
            "الفحوص اليدوية تثبت قابلية البناء والحماية الأساسية فقط، ولا تثبت منطق الأعمال.",
        ],
        "todo": [
            "Unit Tests للرموز والتصحيح والأثر.",
            "Integration Tests لقاعدة البيانات وJotform وxAPI.",
            "RLS Tests لكل دور وجدول، وIdempotency للمصادر والشهادات.",
            "E2E كامل من التسجيل إلى الشهادة، وحالات النقص والفشل والإلغاء.",
            "اختبارات RTL والأجهزة والوصول والأداء وUAT وإغلاق العيوب.",
        ],
        "evidence": "package.json لا يحتوي test أو lint؛ عدد ملفات الاختبار صفر؛ نتائج الفحص التنفيذي الحالية",
        "acceptance": "غير مغلقة؛ لا توجد رحلة E2E أو اختبارات RLS أو أجهزة.",
    },
    {
        "n": 17,
        "title": "ترحيل البيانات والاستعداد للتجربة",
        "status": "لم يبدأ",
        "tone": "todo",
        "summary": "الأرقام التاريخية موجودة في الوثائق والنموذج، ولا يوجد دليل على استيراد 12 قبلي و7 بعدي إلى القاعدة الحالية.",
        "done": [
            "تثبيت الأرقام المرجعية للدفعة في الوثائق والنموذج: 12 قبلي، 7 بعدي، min 5→8، وSD 1.62→0.79.",
        ],
        "partial": [
            "الصفحات التجريبية تعرض هذه الأرقام، وهذا ليس إثبات ترحيل.",
        ],
        "todo": [
            "نسخة محفوظة من مصدر Jotform واستيراد يحفظ submission_id والمصدر الأصلي.",
            "توثيق السجلات غير القابلة للمطابقة دون اختلاق روابط.",
            "مطابقة المتوسط والحد الأدنى والانحراف مع المصدر.",
            "إعداد جهة وبرنامج ودفعة ونماذج ومفاتيح وبطاقات للتجربة.",
            "Dry Run داخلي وخطة دعم وانقطاع إنترنت وتدريب المشرفين.",
        ],
        "evidence": "al-miqyas-v2.jsx، 01-project-brief.md، عدم وجود سكربت ترحيل أو Seed تاريخي",
        "acceptance": "غير مغلقة؛ لا توجد بيانات تاريخية مثبتة في التدفق الحالي ولا Dry Run.",
    },
    {
        "n": 18,
        "title": "النشر الإنتاجي",
        "status": "غير مثبت",
        "tone": "todo",
        "summary": "لا توجد أدلة محلية على Vercel أو الدومين أو SSL أو Tag أو Release. البيئة الحالية تعمل محليًا فقط في نطاق هذا التدقيق.",
        "done": [
            "البناء الإنتاجي المحلي ينجح.",
            "تهيئة cpus=1 تمنع استهلاك موارد زائد أثناء البناء.",
        ],
        "partial": [
            "ملف التسليم يقترح Vercel والدومين، لكن لا توجد vercel.json أو بيانات نشر أو فحص حي.",
            "مهاجرات القاعدة مطبقة وفق التسليم، لكن لم تُربط بإصدار Git موثق.",
        ],
        "todo": [
            "إنشاء Production وتطبيق المهاجرات المعتمدة فقط وإضافة الأسرار في المنصات.",
            "ضبط DNS وSSL وRedirects وتعطيل البيانات التجريبية.",
            "Backup ومراقبة وتنبيهات وFreeze وTag وRelease Notes.",
            "Smoke Tests كاملة على الدومين وخطة Rollback مع مسؤول مخول.",
        ],
        "evidence": "نتيجة build المحلية، غياب vercel.json وCI وGit صالح، وثائق التسليم فقط",
        "acceptance": "غير مغلقة؛ لا يوجد إصدار إنتاجي مثبت أو Smoke Test على دومين.",
    },
    {
        "n": 19,
        "title": "التجربة الميدانية وقياس النجاح",
        "status": "لم تبدأ",
        "tone": "todo",
        "summary": "لا يمكن تنفيذ Pilot قبل Jotform وxAPI والمحرك والشهادات والاختبارات.",
        "done": [
            "مؤشرات النجاح المستهدفة مذكورة في قائمة المهام والوثائق.",
        ],
        "partial": [
            "السيناريو التعليمي جاهز نسبيًا، لكنه لا يساوي جاهزية المنظومة التقنية.",
        ],
        "todo": [
            "تشغيل دورة كاملة ومراقبة القبلي وxAPI والبعدي والتقرير والشهادات.",
            "قياس المطابقة 98% وزمن التقرير أقل من دقيقة وعدم التسرب.",
            "اختبار QR شهادات على iPhone ومقارنة عينة بالمصدر.",
            "تقرير Pilot بالأعطال والتدخلات والقرارات والإصلاحات.",
        ],
        "evidence": "سيناريو تهيئة الموظفين الجدد، قائمة المهام، غياب أي تقرير Pilot",
        "acceptance": "غير مغلقة؛ لا توجد دورة ميدانية كاملة أو مؤشرات مقاسة.",
    },
    {
        "n": 20,
        "title": "التوثيق والتسليم ونقل الملكية",
        "status": "جزئي",
        "tone": "partial",
        "summary": "التوثيق المفاهيمي والتسليم المرحلي جيد، لكن وثائق المطور والتشغيل والأمان والملكية النهائية ناقصة.",
        "done": [
            "دليل تعليمي وتقني، قائمة مهام، وثائق brief وarchitecture وroadmap وintegrations.",
            "سجل قرارات تصميم، ملف تصميم، وملفا تسليم بتاريخي 21 و23 يوليو.",
            "سبع مهاجرات مرتبة ومشروحة داخل SQL.",
        ],
        "partial": [
            "ملف HANDOVER_2026-07-23.md دقيق في فصل الحقيقي عن التجريبي، لكنه لقطة انتقالية لا دليل تشغيل نهائي.",
            "حزمة al-miqyas-handoff قديمة جزئيًا وتحتوي Schema وسياسات superseded.",
        ],
        "todo": [
            "README داخل التطبيق، ERD، Data Dictionary، وتوثيق APIs الحقيقية.",
            "توثيق قواعد الحساب والإصدارات وأمثلة الاختبار.",
            "دليل مدير الجهة والمدرب والشهادات وإطلاق دورة جديدة.",
            "Runbooks للأعطال وإعادة المعالجة والمفاتيح والنسخ والاستعادة.",
            "إثبات ملكية الحسابات وقائمة وصول آمنة وAcceptance Checklist وفترة دعم.",
        ],
        "evidence": "ملفات التحضير، al-miqyas-handoff/docs، ChatGPT docs، HANDOVER_2026-07-21 و2026-07-23",
        "acceptance": "غير مغلقة؛ فريق التشغيل لا يملك بعد أدلة كافية للنشر والاستعادة والتشغيل المستقل.",
    },
    {
        "n": 21,
        "title": "ما بعد الإطلاق والـBacklog",
        "status": "مؤجل كما ينبغي",
        "tone": "note",
        "summary": "هذه المهمة ليست متأخرة؛ هي مؤجلة منطقيًا حتى استقرار MVP. الخطر هو البدء بها قبل إغلاق المهام 0 إلى 20.",
        "done": [
            "تحديد عناصر ما بعد الإطلاق في القائمة الأصلية: بنوك أسئلة بالذكاء الاصطناعي، نقل Jotform، Benchmarking، PDF مخصص، LRS، وتطبيق جوال.",
        ],
        "partial": [
            "لا توجد آلية Backlog أو ترتيب قيمة ومخاطر مثبتة.",
        ],
        "todo": [
            "عدم تنفيذ أي امتداد قبل Pilot ناجح واستقرار الإنتاج.",
            "بعد الإطلاق: مراقبة التكاليف والمطابقة وزمن التقرير وفشل التكاملات.",
            "Retrospective وترتيب Backlog بمعيار قيمة ومخاطر ومعيار قبول.",
        ],
        "evidence": "04-roadmap.md وقائمة المهام الأصلية",
        "acceptance": "مؤجلة؛ لا تُفتح قبل إغلاق MVP وقياسه.",
    },
]


def add_metric_cards(doc):
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    metrics = [
        ("0", "مهمة أساسية مغلقة بالكامل", RED, RED_SOFT),
        ("18", "جدولًا مع RLS", PURPLE, PURPLE_SOFT),
        ("34", "سياسة RLS", GREEN, GREEN_SOFT),
        ("0", "ملف اختبار آلي", RED, RED_SOFT),
        ("15", "مسارًا بُنيت بنجاح", BLUE, BLUE_SOFT),
        ("7", "مهاجرات حالية", PURPLE, PURPLE_SOFT),
        ("0", "مسار app/api", RED, RED_SOFT),
        ("91", "ملفًا دلاليًا في نطاق الجرد", GRAY_700, GRAY_100),
    ]
    for idx, (value, label, accent, fill) in enumerate(metrics):
        r, c = divmod(idx, 4)
        cell = table.cell(r, c)
        set_cell_shading(cell, fill)
        set_cell_border(cell, top={"val": "single", "sz": "6", "color": accent},
                        bottom={"val": "single", "sz": "6", "color": accent},
                        start={"val": "single", "sz": "6", "color": accent},
                        end={"val": "single", "sz": "6", "color": accent})
        set_cell_margins(cell, 130, 110, 130, 110)
        p = cell.paragraphs[0]
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, value, "Noto Sans Arabic", 18, True, accent)
        p2 = cell.add_paragraph()
        set_rtl(p2, WD_ALIGN_PARAGRAPH.CENTER)
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, rt(label), "Noto Sans Arabic", 7.7, False, GRAY_700)
    doc.add_paragraph()


def add_executive_summary(doc):
    add_section_title(doc, "1 · الخلاصة التنفيذية", "النتيجة الصريحة قبل التفاصيل")
    add_callout(
        doc,
        "الحكم",
        "حسب معايير الإغلاق المكتوبة في قائمة المهام، لا توجد مهمة أساسية من 0 إلى 20 مكتملة بالكامل. هذا لا يعني أن المشروع صفر؛ بل يعني أن العمل الحقيقي الحالي يتركز في قاعدة البيانات وتدفقات البرامج والمتدربين، بينما التكامليات والمحرك والتقارير والشهادات والاختبارات والنشر ما زالت مفتوحة.",
        fill=RED_SOFT,
        accent=RED,
    )
    add_metric_cards(doc)
    add_subhead(doc, "ما هو حقيقي الآن")
    for item in [
        "مصادقة Supabase بكلمة مرور، حماية المسارات، وتسجيل خروج.",
        "قاعدة بيانات إنتاجية التصميم تضم 18 جدولًا وقيودًا مركبة وRLS.",
        "إنشاء وتعديل برامج ودفعات مسودة عبر RPCs ذرية مع Audit Log.",
        "إنشاء متدرب وتسجيله وتعديل ملفه، ومعرف AMD-XXXXX موحد.",
        "صفحة متدرب داخلية تقرأ الأدلة الفعلية فقط، وQR فعلي.",
        "صفحة /t/[traineeCode] عامة تقرأ حالة فعلية دون بيانات شخصية.",
    ]:
        add_bullet(doc, item, tone="done")
    add_subhead(doc, "ما هو محاكاة أو غير موجود")
    for item in [
        "لوحة التحكم والجلسات الحية والتقارير والشهادات والتحقق العام وإدارة الجهات والإعدادات والحساب وغرفة الدفعة.",
        "Jotform Webhook وReconciliation والتصحيح.",
        "POST /api/xapi وإدارة مفاتيح الجهات والبث الحي.",
        "محرك القياس والأثر وإصدار الشهادة.",
        "Unit وIntegration وRLS وE2E وCI.",
        "دومين إنتاجي موثق وPilot ونسخ احتياطي مجرب وPDPL وRunbooks.",
    ]:
        add_bullet(doc, item, tone="todo")


def add_methodology(doc):
    add_section_title(doc, "2 · نطاق المراجعة وطريقتها", "ما قُرئ، وما استُبعد، وما تم التحقق منه")
    add_subhead(doc, "نطاق الملفات")
    for item in [
        "الوثيقتان المرجعيتان PDF وDOCX: قائمة مهام التنفيذ، والدليل التعليمي والتقني.",
        "سيناريو تهيئة الموظفين الجدد كاملًا.",
        "حزمة التسليم al-miqyas-handoff ووثائق brief وarchitecture وschema وroadmap وintegrations والنموذج الأولي.",
        "تطبيق al-miqyas الحالي: المسارات والمكونات والإعدادات والأصول وملفات البيئة بأسماء المفاتيح فقط.",
        "جميع المهاجرات الحالية 001 إلى 007 وسياسات RLS وRPCs.",
        "ملفات DESIGN وPROJECT_DECISIONS وWRITING_DIRECTION_GUIDE وملفات HANDOVER.",
    ]:
        add_bullet(doc, item, tone="note")
    add_subhead(doc, "الاستبعادات المنهجية")
    add_paragraph(
        doc,
        "لم تُعامل node_modules و.next وملفات البناء وtsconfig.tsbuildinfo وحزم الطرف الثالث الكبيرة داخل ChatGPT كتنفيذ للمنتج. تم جردها وتصنيفها كمخرجات مولدة أو مكتبات خارجية. كما عوملت النسخة ZIP والنسخة المكررة من al-miqyas-v2.jsx كأرشيفات مطابقة، لا كتنفيذ إضافي.",
        size=9.4,
        color=GRAY_700,
    )
    add_subhead(doc, "التحقق التنفيذي بتاريخ 24 يوليو 2026")
    results = [
        ("typecheck", "نجح", "tsc --noEmit", GREEN),
        ("build", "نجح", "Next.js production build، 15 مسارًا", GREEN),
        ("/login", "200", "صفحة عامة سليمة", GREEN),
        ("/dashboard", "307", "إعادة توجيه صحيحة إلى login عند عدم المصادقة", GREEN),
        ("/t/AMD-AAAAA", "200", "المسار العام يعمل ويعرض حالة عدم العثور", GREEN),
        ("/verify/TEST", "200", "يعمل، لكنه صفحة محاكاة", AMBER),
        ("Supabase RPC", "200", "get_public_trainee_route موجودة فعليًا وأعادت [] لرمز وهمي", GREEN),
        ("Supabase OpenAPI", "401", "لم يُسمح بجرد المخطط عبر المفتاح العام", GRAY_700),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.0), Cm(2.2), Cm(7.2), Cm(4.0)]
    headers = ["الفحص", "النتيجة", "الدليل", "الحكم"]
    for i, h in enumerate(headers):
        table.columns[i].width = widths[i]
        set_cell_text(table.cell(0, i), h, size=8.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for name, result, evidence, color in results:
        row = table.add_row()
        set_cell_text(row.cells[0], name, size=8.2)
        set_cell_text(row.cells[1], result, size=8.2, bold=True, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], evidence, size=8.2)
        judgment = "مثبت" if color == GREEN else ("جزئي" if color == AMBER else "قيد الصلاحية")
        set_cell_text(row.cells[3], judgment, size=8.2, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_borders(table)
    add_callout(
        doc,
        "حدود الإثبات",
        "لم تُستخدم بيانات اعتماد مستخدم حقيقي، ولم تُنفذ أي كتابة في Supabase أثناء التدقيق. لذلك تم إثبات السطح العام والبناء والحماية الأساسية، لكن لا يمكن اعتبار بيانات الإنتاج أو ملكية الحسابات أو عزل جهتين مختبرًا حاليًا.",
        fill=AMBER_SOFT,
        accent=AMBER,
    )


def status_color(tone):
    return {"done": GREEN, "partial": AMBER, "todo": RED, "note": BLUE}[tone]


def status_fill(tone):
    return {"done": GREEN_SOFT, "partial": AMBER_SOFT, "todo": RED_SOFT, "note": BLUE_SOFT}[tone]


def add_master_map(doc):
    add_section_title(doc, "3 · الخريطة العامة المحدثة", "الحالة وفق معيار الإغلاق الأصلي، لا وفق الانطباع البصري")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["رقم", "المهمة", "الحالة", "نوع الموجود", "قرار الإغلاق"]
    widths = [Cm(1.0), Cm(6.4), Cm(3.3), Cm(4.4), Cm(3.1)]
    for i, h in enumerate(headers):
        table.columns[i].width = widths[i]
        set_cell_text(table.cell(0, i), h, size=8.1, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for task in TASKS:
        row = table.add_row()
        tone = task["tone"]
        set_cell_text(row.cells[0], str(task["n"]), size=8.1, bold=True, color=status_color(tone), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], task["title"], size=8.1, bold=True)
        set_cell_text(row.cells[2], task["status"], size=7.8, bold=True, color=status_color(tone), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(row.cells[2], status_fill(tone))
        if tone == "partial":
            kind = "كود/قاعدة حقيقية جزئيًا"
        elif tone == "todo":
            kind = "هيكل أو محاكاة أو غياب"
        else:
            kind = "مؤجل بعد MVP"
        set_cell_text(row.cells[3], kind, size=7.7)
        close = "غير مغلق" if task["n"] != 21 else "مؤجل"
        set_cell_text(row.cells[4], close, size=7.8, bold=True, color=status_color(tone), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_borders(table)
    add_callout(
        doc,
        "قراءة الجدول",
        "وجود جداول أو واجهة لا يساوي اكتمال المهمة. اعتُمد الإغلاق فقط عندما تتحقق جميع عناصر المهمة ومعيار القبول والاختبار الوارد في الوثيقة الأصلية.",
        fill=GRAY_100,
        accent=GRAY_700,
    )


def add_real_vs_demo(doc):
    add_section_title(doc, "4 · خريطة الحقيقي مقابل المحاكاة", "المسارات التي يجوز الاعتماد عليها، والمسارات التي لا يجوز عرضها كإنتاج")
    rows = [
        ("/login", "حقيقي", "Supabase signInWithPassword", "لا يزال قرار OTP متعارضًا"),
        ("/programs", "حقيقي جزئي", "قراءة وإنشاء وتعديل برامج ودفعات مسودة", "لا نشر ولا تكاملات"),
        ("/trainees", "حقيقي جزئي", "CRUD متدرب وتسجيل في دفعة", "لا CSV ولا موافقة PDPL"),
        ("/trainees/[code]", "حقيقي جزئي", "قراءة الأدلة وQR", "لا محرك أثر"),
        ("/t/[traineeCode]", "حقيقي جزئي", "RPC عامة وتوجيه حسب الحالة", "مشكلة تعدد التسجيلات وRate Limiting"),
        ("/dashboard", "محاكاة", "أرقام وأسماء ثابتة", "لا يجوز اعتباره Dashboard فعليًا"),
        ("/sessions", "محاكاة", "أحداث xAPI ثابتة", "لا endpoint ولا بث"),
        ("/reports", "محاكاة", "قيم تاريخية ثابتة", "لا حساب ولا إصدار"),
        ("/certificates", "محاكاة", "إصدار بصري فقط", "لا إصدار أو إلغاء حقيقي"),
        ("/verify/[verifyCode]", "محاكاة", "حالات ثابتة valid/revoked/notfound", "لا قراءة من certificates"),
        ("/organizations", "محاكاة", "جهات وأعضاء ثابتون", "لا إدارة جهة أو دعوات"),
        ("/settings و/account", "محاكاة", "إعدادات وهوية مستخدم ثابتة", "تتضمن نصًا قديمًا يناقض الاتصال الحالي"),
        ("/cohorts/[id]/run", "محاكاة", "12 مشاركًا وأحداث محلية", "لا دفعة فعلية"),
        ("/register", "محاكاة", "استلام طلب تجريبي", "لا حساب أو دعوة"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.1), Cm(2.7), Cm(6.3), Cm(5.1)]
    for i, h in enumerate(["المسار", "الحالة", "ما يفعله", "الفجوة"]):
        table.columns[i].width = widths[i]
        set_cell_text(table.cell(0, i), h, size=8.2, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for route, state, action, gap in rows:
        row = table.add_row()
        color = GREEN if state == "حقيقي" else (AMBER if "جزئي" in state else RED)
        fill = GREEN_SOFT if state == "حقيقي" else (AMBER_SOFT if "جزئي" in state else RED_SOFT)
        set_cell_text(row.cells[0], route, size=7.9)
        set_cell_text(row.cells[1], state, size=7.9, bold=True, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(row.cells[1], fill)
        set_cell_text(row.cells[2], action, size=7.9)
        set_cell_text(row.cells[3], gap, size=7.9)
    set_table_borders(table)
    add_callout(
        doc,
        "مخاطرة إطلاق عالية",
        "الغلاف الحالي يجمع صفحات حقيقية مع صفحات محاكاة تحت نفس التنقل وبنفس الهوية. إذا عُرض هذا على عميل أو استخدم في Pilot دون وسم صارم، فسيفهم المستخدم أن البيانات الحية والتقارير والشهادات تعمل بينما هي لا تعمل.",
        fill=RED_SOFT,
        accent=RED,
    )


def add_database_review(doc):
    add_section_title(doc, "5 · مراجعة القاعدة والكود", "نقاط القوة والعيوب التي تؤثر في التنفيذ")
    add_subhead(doc, "قاعدة البيانات")
    for item in [
        "الانتقال من Schema الحزمة القديمة إلى migrations الحالية أصلح جداول cohorts وenrollments وcohort_reports وwebhook_ingestions وstatement_id والقيود المركبة.",
        "العزل عبر org_id مدعوم بمفاتيح خارجية مركبة، وهذا أقوى من الاعتماد على RLS وحدها.",
        "RLS مفعلة على جميع الجداول الثمانية عشر، والسياسات الحالية لا تمنح viewer كتابة.",
        "خامات التكامل والأثر والشهادات صُممت بإصدارات وIdempotency وsnapshots مناسبة.",
    ]:
        add_bullet(doc, item, tone="done")
    add_subhead(doc, "عيوب يجب عدم تجاهلها")
    for item in [
        "لا توجد اختبارات RLS. التصميم وحده لا يثبت عدم التسرب.",
        "الدالة العامة للمتدرب تختار أحدث تسجيل للرمز، ما قد يوجه خطأ عند تسجيل المتدرب في أكثر من برنامج.",
        "لا توجد API routes، لذلك لا يوجد مكان آمن لتشغيل Jotform أو xAPI أو إصدار الشهادات.",
        "لا يوجد Server Supabase module مستقل أو service-role path، باستثناء عميل جلسة داخل proxy.",
        "الواجهة تقرأ وتكتب مباشرة من المتصفح عبر RLS وRPCs؛ هذا مقبول لبعض CRUD، لكنه غير مناسب للتكاملات والأسرار والمحرك.",
        "لا توجد آلية نشر نسخة برنامج أو ربط نماذج أو مفاتيح من الواجهة.",
    ]:
        add_bullet(doc, item, tone="todo")
    add_subhead(doc, "تناقضات الوثائق والكود")
    contradictions = [
        ("Auth", "الوثائق تطلب OTP", "الكود يستخدم كلمة مرور", "حسم قرار واحد وتحديث الجميع"),
        ("البيئات", "القائمة تطلب Dev/Staging/Production", "قرار 007 يعتمد بيئة واحدة", "استثناء معتمد أو إعادة القرار"),
        ("مفتاح Supabase", "env.example يستخدم ANON_KEY", "التطبيق يستخدم PUBLISHABLE_KEY", "env.example حالي داخل التطبيق"),
        ("الشعار", "brand-spec يقول pending", "أصول الشعار مستخدمة فعليًا", "تحديث brand-spec"),
        ("الإعدادات", "واجهة تقول Supabase غير متصل", "المسارات الحقيقية متصلة", "إزالة النص التجريبي"),
        ("Git", "المهمة تشترط ملكية الأمد", "المجلد الحالي ليس مستودع Git صالحًا", "ربط بمستودع واضح وإثبات الملكية"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(["الموضوع", "المرجع", "الواقع", "المطلوب"]):
        set_cell_text(table.cell(0, i), h, size=8.2, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for row_data in contradictions:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=7.9, bold=(i == 0), color=(RED if i == 2 else GRAY_900))
    set_table_borders(table)


def add_task_detail(doc, task):
    tone = task["tone"]
    color = status_color(tone)
    fill = status_fill(tone)
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_run(p, rt(f"{task['n']} · {task['title']}"), "Noto Kufi Arabic", 15, True, NAVY)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    data = [
        ("الحالة", task["status"]),
        ("الإغلاق", "غير مغلق" if task["n"] != 21 else "مؤجل"),
    ]
    for idx, (label, value) in enumerate(data):
        label_cell = table.cell(0, idx * 2 + 1)
        value_cell = table.cell(0, idx * 2)
        set_cell_text(label_cell, label, size=8.4, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(label_cell, NAVY)
        set_cell_text(value_cell, value, size=8.4, bold=True, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(value_cell, fill)
    set_table_borders(table)

    add_paragraph(doc, task["summary"], size=9.5, color=GRAY_700, before=5, after=6)
    add_subhead(doc, "تم إنجازه فعليًا", GREEN)
    for item in task["done"]:
        add_bullet(doc, item, tone="done")
    if task["partial"]:
        add_subhead(doc, "منجز جزئيًا أو يحتاج إثباتًا", AMBER)
        for item in task["partial"]:
            add_bullet(doc, item, tone="partial")
    add_subhead(doc, "المتبقي للإغلاق", RED)
    for item in task["todo"]:
        add_bullet(doc, item, tone="todo")
    add_subhead(doc, "الدليل")
    add_paragraph(doc, task["evidence"], size=8.7, color=GRAY_700, after=4)
    add_callout(doc, "نتيجة معيار القبول", task["acceptance"], fill=fill, accent=color)


def add_tasks_section(doc):
    add_section_title(doc, "6 · المراجعة التفصيلية للمهام 0 إلى 21", "لكل مهمة: المنجز، الجزئي، المتبقي، والدليل")
    add_paragraph(
        doc,
        "المراجعة التالية لا تعيد كتابة قائمة المهام حرفيًا. هي تحديث تنفيذي لها: كل مجموعة فرعية عولجت ضمن تم أو جزئي أو متبقٍ، ثم قورنت ببوابة القبول الأصلية.",
        size=9.4,
        color=GRAY_700,
    )
    doc.add_page_break()
    for idx, task in enumerate(TASKS):
        add_task_detail(doc, task)
        if idx != len(TASKS) - 1:
            doc.add_page_break()


def add_risks(doc):
    doc.add_page_break()
    add_section_title(doc, "7 · سجل المخاطر الحالي", "المخاطر التي تمنع الإطلاق أو تفسد الثقة في النتائج")
    risks = [
        ("حرج", "خلط الحقيقي بالمحاكاة", "قد يظن العميل أن التقارير والشهادات والجلسات تعمل", "فصل أو وسم المحاكاة فورًا"),
        ("حرج", "لا توجد رحلة E2E", "لا يمكن إثبات أن القبلي واللحظي والبعدي ينتجون تقريرًا وشهادة", "إكمال Jotform ثم xAPI ثم المحرك والاختبارات"),
        ("عالٍ", "غياب اختبارات RLS", "خطأ سياسة واحدة قد يكشف بيانات جهة", "اختبارات حسابين وجهتين لكل جدول"),
        ("عالٍ", "التوجيه يعتمد أحدث تسجيل", "قد يوجه المتدرب إلى برنامج خاطئ", "تمرير سياق enrollment أو route token"),
        ("عالٍ", "لا Rate Limiting للمسار العام", "تعداد رموز AMD-XXXXX واكتشاف حالات تشغيل", "طبقة API وحدود ومراقبة"),
        ("عالٍ", "لا PDPL أو استعادة مجربة", "مخاطر امتثال وفقد بيانات", "اعتماد سياسة واختبار Restore"),
        ("عالٍ", "لا مستودع Git مثبت ولا CI", "لا أثر موثوق للإصدارات أو الحماية من الدمج الخاطئ", "مستودع مملوك وحماية main"),
        ("متوسط", "نسخ قديمة ومتكررة", "خطر تعديل الملف الخطأ وعودة بيانات تجريبية", "تنظيف بعد تأكيد عدم الاستخدام"),
        ("متوسط", "تعارض الوثائق", "المطور ينفذ OTP أو بيئات أو مفاتيح باسم خاطئ", "تحديث وثيقة واحدة مرجعية"),
        ("متوسط", "لا NEXT_PUBLIC_APP_URL", "QR يعمل على الجهاز نفسه ويفشل من الهاتف", "ضبط URL منشور واختبار حقيقي"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(["الشدة", "الخطر", "الأثر", "الإجراء"]):
        set_cell_text(table.cell(0, i), h, size=8.2, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for severity, risk, impact, action in risks:
        row = table.add_row()
        color = RED if severity == "حرج" else (AMBER if severity == "عالٍ" else BLUE)
        fill = RED_SOFT if severity == "حرج" else (AMBER_SOFT if severity == "عالٍ" else BLUE_SOFT)
        set_cell_text(row.cells[0], severity, size=8, bold=True, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(row.cells[0], fill)
        set_cell_text(row.cells[1], risk, size=8, bold=True)
        set_cell_text(row.cells[2], impact, size=8)
        set_cell_text(row.cells[3], action, size=8)
    set_table_borders(table)


def add_execution_plan(doc):
    add_section_title(doc, "8 · خطة التنفيذ من نقطة التوقف الصحيحة", "الترتيب الذي يغلق المخاطر بدل فتح واجهات إضافية")
    gates = [
        (
            "البوابة 1 · حسم Jotform",
            "الآن",
            [
                "اعتماد النموذج القبلي والبعدي وحقل traineeId المقفل.",
                "اعتماد answer_key ونوع الأسئلة والثقة وربطهما بإصدار البرنامج.",
                "عدم تشغيل migration 008 قبل نجاح هذا الاعتماد.",
            ],
            "نماذج وروابط prefill مجربة من الجوال، ومخطط mapping معتمد.",
        ),
        (
            "البوابة 2 · استقبال Jotform",
            "بعد البوابة 1",
            [
                "تطبيق migration 008 فقط، ثم انتظار تأكيد نجاحها قبل أي مهاجرة لاحقة.",
                "تنفيذ Next.js Route Handler موحد للـWebhook وReconciliation.",
                "Idempotency وتسجيل الفشل واختبارات إرسال فعلي.",
            ],
            "إرسال قبلي وبعدي يظهران في Supabase مصححين بلا تكرار.",
        ),
        (
            "البوابة 3 · xAPI",
            "بعد ثبات Jotform",
            [
                "اعتماد القاموس مع AmadXR.",
                "إنشاء POST /api/xapi ومفاتيح الجهات والتحقق والحدود.",
                "اختبار جلسة فعلية، التكرار، الانقطاع، والمفتاح الملغى.",
            ],
            "جلسة كاملة مرتبطة بالتسجيل الصحيح وأحداثها غير مكررة.",
        ),
        (
            "البوابة 4 · المحرك والتقرير والشهادة",
            "بعد اكتمال المصادر",
            [
                "scoring والأثر الفردي والجماعي مع snapshots وإصدارات.",
                "إصدار شهادة عند البعدي 80% أو أكثر، وإلغاء وإعادة إصدار.",
                "تحقق عام فعلي وQR مطبوع.",
            ],
            "رحلة بيانات كاملة قابلة لإعادة الإنتاج والتدقيق.",
        ),
        (
            "البوابة 5 · استبدال المحاكاة",
            "قبل أي عرض إنتاجي",
            [
                "ربط dashboard والجلسات والتقارير والشهادات والجهات والإعدادات.",
                "حذف أو عزل الملفات التجريبية والبيانات الثابتة.",
                "ربط مبدل الجهة والهوية والمستخدم بالبيانات الفعلية.",
            ],
            "لا تعرض أي صفحة تشغيلية بيانات غير حقيقية دون وسم.",
        ),
        (
            "البوابة 6 · الجودة والأمن والإطلاق",
            "قبل Pilot",
            [
                "Unit وIntegration وRLS وE2E والأجهزة والوصول.",
                "PDPL وRate Limiting وCSP وBackup/Restore وRunbooks.",
                "Staging أو استثناء موثق، ثم Pilot، ثم Production وSmoke Tests.",
            ],
            "بوابات المهام 0 إلى 20 مغلقة أو لها استثناءات مكتوبة معتمدة.",
        ),
    ]
    for idx, (title, when, actions, gate) in enumerate(gates, start=1):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(15.6)
        table.columns[1].width = Cm(2.2)
        set_cell_text(table.cell(0, 0), title, size=10, bold=True, color=NAVY)
        set_cell_text(table.cell(0, 1), when, size=8, bold=True, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, 1), PURPLE_SOFT)
        set_table_borders(table)
        for action in actions:
            add_bullet(doc, action, tone="todo")
        add_callout(doc, "بوابة الانتقال", gate, fill=GREEN_SOFT, accent=GREEN)
    add_callout(
        doc,
        "نقطة الاستئناف الدقيقة",
        "ابدأ باعتماد نماذج Jotform القبلي والبعدي وحقل traineeId ومفتاح الإجابة. بعدها فقط تُنفذ migration 008، مهاجرة واحدة، ثم يُنتظر تأكيد النجاح قبل الانتقال إلى xAPI.",
        fill=PURPLE_SOFT,
        accent=PURPLE,
    )


def add_file_scope(doc):
    add_section_title(doc, "9 · سجل الملفات والنطاق", "ما اعتُبر مصدرًا، وما اعتُبر تنفيذًا، وما استُبعد")
    categories = [
        ("متطلبات ومرجع", "قائمة المهام PDF/DOCX، الدليل PDF/DOCX، سيناريو التدريب، README ووثائق brief وarchitecture وroadmap وintegrations."),
        ("قرارات وتصميم", "DESIGN.md وPROJECT_DECISIONS.md وWRITING_DIRECTION_GUIDE.md وbrand-spec.md."),
        ("تطبيق حالي", "app وcomponents وlib وproxy وإعدادات Next.js وTypeScript وTailwind والحزم."),
        ("قاعدة بيانات", "سبع مهاجرات SQL حالية 001 إلى 007."),
        ("أصول", "شعارات الجذر وأصول public/brand وapp/icon.png."),
        ("أرشيف", "al-miqyas-handoff.zip والنموذج al-miqyas-v2.jsx المكرر المطابق."),
        ("نسخ قديمة داخل التطبيق", "globals 2.css وicons 2.tsx وtrainee-workspace.tsx وtrainee-workspace 2.tsx وأجزاء قديمة داخل dashboard-pages.tsx."),
        ("مستبعد من تقييم المنتج", "node_modules و.next وملفات البناء وtsconfig.tsbuildinfo ومستودعا Frontend Design Deslop وWeb Design Engineer ومراجع de-slop العامة."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.1)
    table.columns[1].width = Cm(14.0)
    for i, h in enumerate(["التصنيف", "المحتوى"]):
        set_cell_text(table.cell(0, i), h, size=8.6, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for category, contents in categories:
        row = table.add_row()
        set_cell_text(row.cells[0], category, size=8.4, bold=True, color=PURPLE)
        set_cell_shading(row.cells[0], PURPLE_SOFT)
        set_cell_text(row.cells[1], contents, size=8.4)
    set_table_borders(table)
    add_subhead(doc, "ملاحظات جودة الملفات")
    for item in [
        "النموذج al-miqyas-v2.jsx في الجذر وداخل الحزمة متطابقان بالـSHA-1؛ هو مرجع بصري لا تنفيذ إنتاجي.",
        "المجلد الحالي ليس مستودع Git صالحًا للمشروع؛ أمر git status التقط نطاقًا أوسع من المشروع، لذلك لم يُستخدم كدليل حالة.",
        "حزمة التسليم القديمة مفيدة تاريخيًا، لكن Schema وسياسات RLS فيها لا تمثل الوضع الحالي.",
        "ملفات التسليم الحالية تذكر بوضوح الصفحات الحقيقية والتجريبية، وهو توثيق صحيح يجب الحفاظ عليه.",
    ]:
        add_bullet(doc, item, tone="note")


def add_final_gate(doc):
    add_section_title(doc, "10 · قرار الإغلاق الحالي", "هل المشروع مكتمل؟")
    add_callout(
        doc,
        "القرار",
        "لا. المشروع غير مكتمل وغير جاهز للإطلاق أو Pilot حقيقي. القاعدة وتدفقات CRUD الأساسية تقدمت بوضوح، لكن الرحلة التي تعطي المنتج قيمته لم تُنفذ بعد.",
        fill=RED_SOFT,
        accent=RED,
    )
    add_subhead(doc, "شروط الإغلاق النهائي")
    conditions = [
        ("جميع المهام 0 إلى 20 مكتملة أو لها استثناء معتمد", "غير متحقق"),
        ("مطابقة 98% أو أعلى في Pilot", "غير مختبر"),
        ("تقرير خلال أقل من دقيقة", "غير مختبر"),
        ("لا تسرب بين جهتين", "مصمم عبر RLS وغير مختبر"),
        ("Jotform وxAPI والتقرير والشهادة تعمل من طرف إلى طرف", "غير متحقق"),
        ("لا عيوب حرجة أو عالية", "لا توجد منظومة عيوب أو QA"),
        ("الملكية والحسابات والوثائق تحت الأمد", "غير مثبت بالكامل"),
        ("مسؤول تشغيل ودعم وخطط استعادة", "غير موثق"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(["الشرط", "الوضع الحالي"]):
        set_cell_text(table.cell(0, i), h, size=8.6, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_header(table.rows[0])
    for condition, state in conditions:
        row = table.add_row()
        set_cell_text(row.cells[0], condition, size=8.4)
        set_cell_text(row.cells[1], state, size=8.4, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(row.cells[1], RED_SOFT)
    set_table_borders(table)
    add_callout(
        doc,
        "الخطوة التالية الوحيدة الصحيحة",
        "إغلاق نماذج Jotform ومفتاح الإجابة، ثم migration 008. فتح تقارير أو شهادات أو لوحة تحكم جديدة قبل ذلك سيزيد الديون ويكرر المحاكاة.",
        fill=PURPLE_SOFT,
        accent=PURPLE,
    )
    add_paragraph(doc, "نهاية التقرير · لقطة 24 يوليو 2026", size=9, color=GRAY_500, before=28, align=WD_ALIGN_PARAGRAPH.CENTER)


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_executive_summary(doc)
    add_methodology(doc)
    add_master_map(doc)
    add_real_vs_demo(doc)
    add_database_review(doc)
    add_tasks_section(doc)
    add_risks(doc)
    add_execution_plan(doc)
    add_file_scope(doc)
    add_final_gate(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
